"""DOCX document extractor"""

from typing import Dict, Any, Tuple
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from src.knowledge.extractors.base_extractor import BaseExtractor
from src.utils.errors import KnowledgeBaseError
from src.utils.logging import get_logger

logger = get_logger(__name__)


class DOCXExtractor(BaseExtractor):
    """Extract text and metadata from DOCX files"""

    def supports(self, file_path: str) -> bool:
        """Check if file is a DOCX"""
        return Path(file_path).suffix.lower() in [".docx", ".doc"]

    def extract(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from DOCX
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (content, metadata)
        """
        if not DOCX_AVAILABLE:
            raise KnowledgeBaseError("python-docx is required for DOCX extraction")
        
        metadata = self._get_file_metadata(file_path)
        content_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "keywords": doc.core_properties.keywords or "",
                    "comments": doc.core_properties.comments or "",
                    "category": doc.core_properties.category or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                })
            
            # Extract paragraphs with structure preservation
            for para in doc.paragraphs:
                if para.text.strip():
                    # Preserve heading styles
                    if para.style.name.startswith("Heading"):
                        level = para.style.name.replace("Heading", "").strip() or "1"
                        content_parts.append(f"\n{'#' * int(level)} {para.text}\n")
                    else:
                        content_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._format_table(table)
                content_parts.append(f"\n[Table]\n{table_text}\n")
            
            content = "\n".join(content_parts)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            return content, metadata
            
        except Exception as e:
            logger.error("docx_extraction_error", error=str(e), file_path=file_path)
            raise KnowledgeBaseError(f"Failed to extract DOCX: {str(e)}") from e

    def _format_table(self, table) -> str:
        """Format table as text"""
        rows = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows.append(row_text)
        return "\n".join(rows)

